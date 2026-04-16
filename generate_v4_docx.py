#!/usr/bin/env python3
"""
Generate Summer Camp 5 สันเล็ก v4.docx
Updates all game stats from v3.0 to v4.0 based on current gameData.ts values.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ===================== HELPERS =====================

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_run(para, text, bold=None, italic=None, size=None, font_name=None, color=None):
    """Add a run with formatting."""
    run = para.add_run(text)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size is not None:
        run.font.size = size
    if font_name is not None:
        run.font.name = font_name
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return run

def set_cell_text(cell, text, bold=None, italic=None, size=Pt(9), font_name='Arial Unicode MS', color='333333', alignment=None):
    """Set cell text with formatting, clearing existing content."""
    cell.text = ''
    para = cell.paragraphs[0]
    if alignment is not None:
        para.alignment = alignment
    run = para.add_run(text)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size is not None:
        run.font.size = size
    if font_name is not None:
        run.font.name = font_name
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return run

def set_cell_multiline(cell, lines, bold_first=True, size=Pt(9), font_name='Arial Unicode MS', color='333333'):
    """Set cell with multiple lines (newline-separated)."""
    cell.text = ''
    para = cell.paragraphs[0]
    for i, line in enumerate(lines):
        if i > 0:
            run = para.add_run('\n')
            run.font.size = size
            run.font.name = font_name
            run.font.color.rgb = RGBColor.from_string(color)
        is_first_part = (i == 0) if bold_first else False
        run = para.add_run(line)
        run.bold = is_first_part if bold_first else None
        run.font.size = size
        run.font.name = font_name
        run.font.color.rgb = RGBColor.from_string(color)

def make_header_cell(cell, text, fill_color, size=Pt(9), font_name='Arial Unicode MS'):
    """Format a header cell with white text on colored background."""
    set_cell_shading(cell, fill_color)
    set_cell_text(cell, text, bold=True, size=size, font_name=font_name, color='FFFFFF')

def make_data_cell(cell, text, bold=None, fill_color='f8f9fa', size=Pt(9), font_name='Arial Unicode MS', color='333333'):
    """Format a data cell with light background."""
    set_cell_shading(cell, fill_color)
    set_cell_text(cell, text, bold=bold, size=size, font_name=font_name, color=color)

def add_monster_table(doc, header_text, header_color, rows_data, num_rows=13):
    """Add a monster stat table."""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    make_header_cell(table.rows[0].cells[0], rows_data[0][0], header_color)
    make_header_cell(table.rows[0].cells[1], rows_data[0][1], header_color)

    # Data rows
    for i in range(1, len(rows_data)):
        left_text, right_text = rows_data[i]
        fill = 'f8f9fa'
        set_cell_shading(table.rows[i].cells[0], fill)
        set_cell_text(table.rows[i].cells[0], left_text, bold=True, size=Pt(9), font_name='Arial Unicode MS', color='333333')
        set_cell_shading(table.rows[i].cells[1], fill)

        # Handle multiline
        if '\n' in right_text:
            lines = right_text.split('\n')
            set_cell_multiline(table.rows[i].cells[1], lines, bold_first=False)
        else:
            set_cell_text(table.rows[i].cells[1], right_text, size=Pt(9), font_name='Arial Unicode MS', color='333333')

    return table

def add_heading1(doc, text):
    para = doc.add_paragraph()
    para.style = doc.styles['Heading 1']
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Arial Unicode MS'
    run.font.color.rgb = RGBColor.from_string('2C3E50')
    return para

def add_heading2(doc, text):
    para = doc.add_paragraph()
    para.style = doc.styles['Heading 2']
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial Unicode MS'
    run.font.color.rgb = RGBColor.from_string('34495E')
    return para

def add_heading3(doc, text):
    para = doc.add_paragraph()
    para.style = doc.styles['Heading 3']
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial Unicode MS'
    run.font.color.rgb = RGBColor.from_string('7F8C8D')
    return para

def add_normal(doc, text, bold=None, italic=None, size=Pt(10), font_name='Arial Unicode MS', color='333333', alignment=None):
    para = doc.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    if text:
        run = para.add_run(text)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        run.font.size = size
        run.font.name = font_name
        run.font.color.rgb = RGBColor.from_string(color)
    return para

def add_bullet(doc, text, bold_prefix=None, size=Pt(10), font_name='Arial Unicode MS', color='333333'):
    """Add a bullet point with optional bold prefix."""
    para = doc.add_paragraph()
    if bold_prefix:
        run = para.add_run(bold_prefix)
        run.bold = True
        run.font.size = size
        run.font.name = font_name
        run.font.color.rgb = RGBColor.from_string(color)
        remaining = text[len(bold_prefix):]
        run = para.add_run(remaining)
        run.font.size = size
        run.font.name = font_name
        run.font.color.rgb = RGBColor.from_string(color)
    else:
        run = para.add_run(text)
        run.font.size = size
        run.font.name = font_name
        run.font.color.rgb = RGBColor.from_string(color)
    return para


# ===================== MAIN DOCUMENT =====================

doc = Document()

# Page setup
for section in doc.sections:
    section.page_width = Emu(7772400)
    section.page_height = Emu(10058400)
    section.top_margin = Emu(457200)
    section.bottom_margin = Emu(457200)
    section.left_margin = Emu(457200)
    section.right_margin = Emu(457200)

# ==================== FRONT MATTER ====================

# Title
p = doc.add_paragraph()
p.style = doc.styles['Title']
run = p.add_run('Theme DnD ')
run = p.add_run('(2Hr)')
run.bold = False

# Subtitle
p = doc.add_paragraph()
p.style = doc.styles['Subtitle']
p.add_run('ผู้ดูแล: พี่เว + พี่กานต์')

doc.add_paragraph()  # spacer

# ==================== เนื้อเรื่อง + การแสดงเปิด ====================
add_heading1(doc, 'เนื้อเรื่อง + การแสดงเปิด')
add_normal(doc, '[TBD - เนื้อเรื่องเปิดกิจกรรม]', italic=True)

doc.add_paragraph(' ')

# ==================== หน้าที่ กำหนดการ ====================
add_heading1(doc, 'หน้าที่ กำหนดการ (19 คน)')

doc.add_paragraph(' ')

# Staff table (Table 0)
table = doc.add_table(rows=7, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['บทบาท', 'จำนวน', 'ชื่อ / หน้าที่']
for ci, h in enumerate(headers):
    make_header_cell(table.rows[0].cells[ci], h, '2c3e50')
staff_data = [
    ['DM (Dungeon Master)', '5 คน', 'พี่โซดา พี่ปาล์ม พี่อิง พี่ดิว ….'],
    ['GM', '1 คน', 'พี่เว'],
    ['ร้านค้า', '1-2 คน', 'ข้าวฟ่าง …. (ร้านค้ารวม)'],
    ['จดคะแนน Boss', '1 คน', 'ชีต้า'],
    ['Monster', '9 คน', 'ชิน พี เม้ก อาร์ม เหนือ กัส [+2]'],
    ['Support', 'ที่เหลือ', 'ที่เหลือทั้งหมด'],
]
for ri, row_data in enumerate(staff_data):
    for ci, text in enumerate(row_data):
        make_data_cell(table.rows[ri+1].cells[ci], text)

# ==================== อุปกรณ์ ====================
add_heading1(doc, 'อุปกรณ์')
items = [
    '1. \tการ์ด Monster (ใบอธิบายมอนฯ สำหรับพี่แต่ละจุด)',
    '2. \tดาบ (สำหรับ Fighter)',
    '3. \tคฑา (สำหรับ Mage)',
    '4. \tสมุดเวทย์ (สำหรับ Priest)',
    '5. \tลูกเต๋าใหญ่ d4 จำนวน 4-5 ลูก -> กัส เหนือทำ หรือซื้อ',
    '6. \tลูกเต๋า d6 อย่างน้อย 5 ชุด (ทีมละ 1 ชุด)',
    '7. \tเงิน LOC (กระดาษ/เหรียญจำลอง)',
    '8. \tการ์ดไอเทม/สกิล สำหรับแจกเมื่อดรอป',
]
for item in items:
    add_normal(doc, item)

# Spacers before cover page
doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')
doc.add_paragraph(' ')
doc.add_paragraph(' ')

# ==================== COVER PAGE ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LEAGUES OF CODE')
run.bold = True
run.font.size = Pt(24)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor.from_string('2C3E50')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Summer Camp 5')
run.font.size = Pt(24)
run.font.name = 'Arial'

doc.add_paragraph(' ')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Dungeons & Dragons')
run.bold = True
run.font.size = Pt(26)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor.from_string('C0392B')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('คู่มือกิจกรรมสันทนาการ')
run.font.size = Pt(16)
run.font.name = 'Arial Unicode MS'
run.font.color.rgb = RGBColor.from_string('E67E22')

doc.add_paragraph(' ')
doc.add_paragraph(' ')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('เอกสารสำหรับทีมจัดงานเท่านั้น')
run.italic = True
run.font.size = Pt(11)
run.font.name = 'Arial Unicode MS'
run.font.color.rgb = RGBColor.from_string('95A5A6')

# CHANGED: v3.0 -> v4.0
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DRAFT v4.0')
run.bold = True
run.font.size = Pt(11)
run.font.name = 'Arial'
run.font.color.rgb = RGBColor.from_string('E74C3C')

doc.add_paragraph(' ')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('เวลาเล่น: ~2 ชั่วโมง | ผู้เล่น: 5 ทีม ทีมละ ~8 คน ')
run.font.size = Pt(9)
run.font.name = 'Arial Unicode MS'
run.font.color.rgb = RGBColor.from_string('7F8C8D')
run = p.add_run('| Staff: ~18 คน')
run.font.size = Pt(9)
run.font.name = 'Arial Unicode MS'
run.font.color.rgb = RGBColor.from_string('7F8C8D')

doc.add_paragraph()
doc.add_paragraph(' ')

# ==================== 1. ภาพรวมกิจกรรม ====================
add_heading1(doc, '1. ภาพรวมกิจกรรม')
add_normal(doc, 'กิจกรรมสันทนาการในธีม Dungeons & Dragons สำหรับค่าย Leagues of Code Summer Camp 5 น้อง ๆ จะแบ่งทีม สำรวจดันเจี้ยน (ตามแผนที่มหาวิทยาลัย) ต่อสู้กับมอนสเตอร์ เก็บของ ซื้อไอเทม และท้าชิงบอส ทีมที่ทำ damage รวมให้บอสได้มากที่สุดจะเป็นผู้ชนะ')

doc.add_paragraph(' ')

# ==================== 2. ลำดับการเล่น (Game Flow) ====================
add_heading1(doc, '2. ลำดับการเล่น (Game Flow)')

doc.add_paragraph(' ')

# Game Flow table (Table 1)
table = doc.add_table(rows=8, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(['ขั้น', 'กิจกรรม', 'รายละเอียด']):
    make_header_cell(table.rows[0].cells[ci], h, '2c3e50')
flow_data = [
    ['1', 'จับทีม', 'แบ่งน้อง ๆ เป็น 5 ทีม ทีมละ ~8 คน แต่ละทีมได้ DM 1 คน'],
    ['2', 'เลือก Class', 'แต่ละทีมแบ่งคนเป็น Fighter, Priest, Mage (คนที่เป็น class เดียวกัน คุมตัวละครเดียวกัน ตัดสินใจร่วมกัน)'],
    ['3', 'อธิบายกฎ', 'DM อธิบายกฎ ระบบลูกเต๋า และแจกใบตัวละคร'],
    ['4', 'สำรวจดันเจี้ยน', 'ทีมเดินตามแผนที่มหาวิทยาลัย เลือกว่าจะไปตีมอนฯ / ซื้อของ / หรือท้าบอส'],
    ['5', 'ต่อสู้ & ซื้อของ', 'ตีมอนฯ เก็บของ ซื้อไอเทม วางแผนเรื่อง damage type'],
    ['6', 'ท้าบอส', 'เมื่อพร้อม ทีมไปท้าบอส (ตีได้หลายรอบ!) จดบันทึก Total Damage'],
    ['7', 'สรุปผล', 'ทีมที่ทำ Total Damage รวมสูงสุดจากบอส ชนะ!'],
]
for ri, row_data in enumerate(flow_data):
    for ci, text in enumerate(row_data):
        make_data_cell(table.rows[ri+1].cells[ci], text)

# ==================== 3. ทีมงาน (Staff Roles) ====================
add_heading1(doc, '3. ทีมงาน (Staff Roles)')

doc.add_paragraph(' ')

add_normal(doc, 'รวมทีมงานที่ต้องการ: ประมาณ 18-19 คน (5 DM + 9 พี่มอนฯ/บอส + 1 บันทึกคะแนน + 2 ร้านค้า + 1 GM + ...)', italic=True)

doc.add_paragraph(' ')

# Staff roles table (Table 2)
table = doc.add_table(rows=8, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(['บทบาท', 'จำนวน', 'หน้าที่']):
    make_header_cell(table.rows[0].cells[ci], h, '2c3e50')
roles_data = [
    ['DM (Dungeon Master)', '5 คน (1/ทีม)', 'เดินไปกับทีม คอยดูแลกฎ ติดตาม HP/ไอเทม/เงิน อำนวยความสะดวกในการต่อสู้'],
    ['พี่มอนสเตอร์', '9 คน', 'นั่งประจำจุด Rank C x3, Rank B x2, Rank A x1, Rank S x2 (Duo คนละตัว), Boss x1 ใช้ใบอธิบายในการต่อสู้'],
    ['พี่บันทึก Damage', '1 คน', 'ประจำจุด Boss จดบันทึก Total Damage ที่แต่ละทีมทำได้'],
    ['พี่ร้านค้า', '1-2 คน', 'ร้านค้ารวม (ทุกไอเทม) ขายของ + Full Heal'],
    ['GM', '1 คน', 'ผู้คุมเกม ดูแลระบบและความเรียบร้อยภาพรวม ตอบคำถามจากทีมงานทุกคน'],
    ['Cameraman', '1 คน', 'ถ่ายภาพกิจกรรม'],
    ['Support', 'ที่เหลือ', 'คอยซัพพอร์ตร้านค้า มอนสเตอร์ ดูแลน้อง ๆ ที่ DM ไม่สามารถดูแลได้'],
]
for ri, row_data in enumerate(roles_data):
    for ci, text in enumerate(row_data):
        make_data_cell(table.rows[ri+1].cells[ci], text)

# ==================== 4. การจัดทีม ====================
add_heading1(doc, '4. การจัดทีม')

# 4.1 การแบ่งทีม
add_heading2(doc, '4.1 การแบ่งทีม')
add_normal(doc, 'น้อง ๆ ถูกแบ่งเป็น 5 ทีม ทีมละ ~8 คน แต่ละทีมมี DM (รุ่นพี่) 1 คนคอยดูแล เดินทางไปด้วยกันตลอด')

doc.add_paragraph(' ')

# 4.2 การเลือก Class
add_heading2(doc, '4.2 การเลือก Class')
add_normal(doc, 'ในแต่ละทีมจะมี 3 ตัวละคร จาก 3 Class: Fighter, Priest, Mage น้อง ๆ ที่เลือก class เดียวกันจะคุมตัวละครเดียวกัน ต้องตัดสินใจร่วมกันว่าจะทำอะไรในแต่ละเทิร์น')
add_normal(doc, 'ตัวอย่าง: ทีมมี 8 คน อาจแบ่งเป็น Fighter 2 คน, Priest 3 คน, Mage 3 คน (แต่ในเกมยังมีแค่ 3 ตัวละคร)', italic=True)

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# ==================== 5. ตัวละคร (Character Classes) ====================
add_heading1(doc, '5. ตัวละคร (Character Classes)')

doc.add_paragraph(' ')

# Skill legend
p = doc.add_paragraph()
run = p.add_run('สกิล (พื้นฐาน) = ใช้ได้ตลอด | (ปลดล็อค) = ต้องหาจาก drop/ร้านค้า ปลดล็อคถาวร | (ใช้ครั้งเดียว) = ต้องหาใหม่หลังใช้')
run.italic = True
run.font.size = Pt(9)
run.font.name = 'Arial Unicode MS'
run.font.color.rgb = RGBColor.from_string('E74C3C')

doc.add_paragraph(' ')

# 5.1 AC
add_heading2(doc, '5.1 AC (Armor Class) ของตัวละคร')
add_normal(doc, 'เมื่อผู้เล่นหรือมอนฯ โจมตี ให้ทอย 2d6 (ลูกเต๋า d6 สองลูก รวมกัน) เทียบกับ AC เป้าหมาย')
add_normal(doc, 'Fighter AC: 7 | Priest AC: 6 | Mage AC: 6', bold=True)

doc.add_paragraph(' ')

# Character table (Table 3) - UPDATED VALUES
table = doc.add_table(rows=4, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_3 = ['Class', 'HP', 'AC', 'Skills', 'อุปกรณ์']
for ci, h in enumerate(headers_3):
    make_header_cell(table.rows[0].cells[ci], h, '2c3e50')

# Fighter - UPDATED: HP 22, AC 7
fighter_skills = (
    'Sword Slash (พื้นฐาน): Physical Damage 1d6+1d4\n'
    'Shield Wall (พื้นฐาน): เพิ่ม AC ของทั้งทีมขึ้นอีก 2 เป็นเวลา 1 เทิร์น\n'
    'Taunt (ปลดล็อค): ล่อการโจมตีมอนฯ มาที่ตัวเอง 1 เทิร์น\n'
    'Second Wind (ปลดล็อค): ฮีลตัวเอง 25% ของ Max HP (ปัดลง)\n'
    'Stun Strike (ใช้ครั้งเดียว): Physical Damage 3d6+1d4 + Stun มอนฯ 1 เทิร์น'
)
fighter_row = ['Fighter', '22', '7', fighter_skills, 'ดาบ\nโล่?\nผ้าคลุม?']

# Priest - UPDATED: HP 18, AC 6, Armor Break changed to -3 AC (not halve)
priest_skills = (
    'Heal (พื้นฐาน): ฮีล HP เพื่อน 1 คน = 1d6+1d4 HP (ไม่ต้อง roll attack)\n'
    'Amplify Damage (พื้นฐาน): บัฟ damage ตัวละคร 1 คน +1d6 ในเทิร์นถัดไป\n'
    'Armor Break (ปลดล็อค): ลด AC มอนฯ -3 เป็นเวลา 1 เทิร์น\n'
    'Daze (ปลดล็อค): มอนฯ ติด Disadvantage 1 เทิร์น\n'
    'Empower (ปลดล็อค): เพื่อน 1 คน ติด Advantage 1 เทิร์น'
)
priest_row = ['Priest', '18', '6', priest_skills, 'หนังสือ/Tome\nมงกุฏสวมหัว']

# Mage - HP 14, AC 6 (same as before)
mage_skills = (
    'Staff Strike (พื้นฐาน): Physical Damage 1d6\n'
    'Fireball (พื้นฐาน): Fire Damage 1d6+1d4\n'
    'Frost Bolt (ปลดล็อค): Frost Damage 1d6+1d4\n'
    'Lightning (ปลดล็อค): Lightning Damage 2d6\n'
    'Necrotic Blast (ปลดล็อค): Necrotic Damage 2d6\n'
    'Holy Light (ปลดล็อค): Holy Damage 3d6'
)
mage_row = ['Mage', '14', '6', mage_skills, 'หมวกแม่มด?\nคฑาเวทย์\nผ้าคลุม']

char_data = [fighter_row, priest_row, mage_row]
class_colors = ['fdebd0', 'fdebd0', 'fdebd0']  # same light orange for all

for ri, row_data in enumerate(char_data):
    fill = 'fdebd0'
    for ci, text in enumerate(row_data):
        cell = table.rows[ri+1].cells[ci]
        set_cell_shading(cell, fill)
        if '\n' in text:
            lines = text.split('\n')
            cell.text = ''
            para = cell.paragraphs[0]
            for li, line in enumerate(lines):
                if li > 0:
                    r = para.add_run('\n')
                    r.font.size = Pt(9)
                    r.font.name = 'Arial Unicode MS' if ci >= 3 else 'Arial'
                    r.font.color.rgb = RGBColor.from_string('333333')
                # Skill names bold
                if ci == 3 and ':' in line:
                    parts = line.split(':', 1)
                    r = para.add_run(parts[0] + ':')
                    r.bold = True
                    r.font.size = Pt(9)
                    r.font.name = 'Arial'
                    r.font.color.rgb = RGBColor.from_string('333333')
                    r = para.add_run(parts[1])
                    r.font.size = Pt(9)
                    r.font.name = 'Arial Unicode MS'
                    r.font.color.rgb = RGBColor.from_string('333333')
                else:
                    fn = 'Arial Unicode MS' if ci >= 3 else 'Arial'
                    r = para.add_run(line)
                    r.font.size = Pt(9) if ci != 4 else Pt(8)
                    r.font.name = fn
                    r.font.color.rgb = RGBColor.from_string('333333')
        else:
            is_bold = (ci == 0)
            fn = 'Arial' if ci < 3 else 'Arial Unicode MS'
            set_cell_text(cell, text, bold=is_bold, size=Pt(9), font_name=fn, color='333333')

# ==================== 6. ระบบลูกเต๋า (Dice Mechanics) ====================
add_heading1(doc, '6. ระบบลูกเต๋า (Dice Mechanics)')
add_normal(doc, 'เกมนี้ใช้ลูกเต๋า 2 ชนิดเท่านั้น: d4 (4 หน้า) และ d6 (6 หน้า) ทอยจริง ๆ ด้วยมือ ระบบ Attack Roll ใช้ 2d6 (ทอย d6 สองลูก รวมกัน ได้ค่า 2-12)')

doc.add_paragraph(' ')

# Dice table (Table 4)
table = doc.add_table(rows=5, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(['ลูกเต๋า', 'หน้า', 'ค่า', 'ใช้ทำอะไร']):
    make_header_cell(table.rows[0].cells[ci], h, '27ae60')
dice_data = [
    ['d4', '4', '1-4', 'Damage ต่ำ / Buff เล็กน้อย'],
    ['d6', '6', '1-6', 'Damage ปานกลาง / Heal / Attack Roll (2d6)'],
    ['2d6', '2 ลูก', '2-12', 'Attack Roll / Initiative (ค่าเฉลี่ย 7)'],
    ['-', '-', '-', '(ไม่ใช้ d8, d20 ในเกมนี้)'],
]
for ri, row_data in enumerate(dice_data):
    for ci, text in enumerate(row_data):
        make_data_cell(table.rows[ri+1].cells[ci], text)

doc.add_paragraph(' ')
doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# ==================== 7. ระบบการต่อสู้ (Combat System) ====================
add_heading1(doc, '7. ระบบการต่อสู้ (Combat System)')

# 7.1 Combat Flow
add_heading2(doc, '7.1 ลำดับการต่อสู้ (Combat Flow)')

doc.add_paragraph(' ')

# Combat flow table (Table 5)
table = doc.add_table(rows=11, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(['ขั้น', 'สิ่งที่ทำ']):
    make_header_cell(table.rows[0].cells[ci], h, 'e67e22')
combat_flow = [
    ['1', 'ทีมเดินทางไปถึงจุดมอนสเตอร์ พี่ที่ประจำจุดทักทาย เริ่มไฟท์ด้วยการใช้อาวุธสกิดพี่มอนฯ'],
    ['2', 'ทุกตัวละคร + มอนฯ ทอย Initiative (2d6) ค่าสูงสุดได้เล่นก่อน เรียงจากมากไปน้อย ถ้าเท่ากัน คนทอยก่อนไปก่อน'],
    ['3', 'เทิร์นของตัวละคร: เลือกใช้สกิล 1 อย่าง + ใช้ไอเทม 1 ชิ้น (ก่อนหรือหลังสกิลก็ได้)'],
    ['4', 'หากเลือกโจมตี: ทอย 2d6 (Attack Roll) ระบบ: 2=Crit Fail (miss) | < AC=Glancing Blow (dmg÷2) | >= AC=Hit (dmg เต็ม) | 12=Crit Hit (dmg×2)'],
    ['5', 'ตีโดน: ทอย Damage Dice ตามสกิล/อาวุธ แล้วคำนวณตาม Vulnerability (x2 / ÷2 / 0) | Glancing Blow: damage ÷ 2 (ปัดลง ขั้นต่ำ 1)'],
    ['6', 'ทอยได้ 2 (Critical Failure): ไม่เกิดอะไร จบเทิร์น (0 damage)'],
    ['7', 'เทิร์นมอนฯ: มอนฯ ทอย 2d6 (Attack Roll) ใช้ระบบเดียวกัน: 2=miss, <AC=dmg÷2, >=AC=full, 12=dmg×2'],
    ['8', 'วนเทิร์นตาม Initiative จนกว่ามอนฯ จะตาย หรือทีมหนี'],
    ['9', 'มอนฯ ตาย: เก็บของดรอป (เงิน LOC / ไอเทม / สกิล) มอนฯ เข้า Spawn Cooldown'],
    ['10', 'เดินทางไปจุดถัดไปบนแผนที่มหาวิทยาลัย'],
]
for ri, row_data in enumerate(combat_flow):
    for ci, text in enumerate(row_data):
        make_data_cell(table.rows[ri+1].cells[ci], text)

# 7.2 Advantage & Disadvantage
add_heading2(doc, '7.2 Advantage & Disadvantage')

doc.add_paragraph(' ')

# Adv/Disadv table (Table 6)
table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(['สถานะ', 'วิธีการ']):
    make_header_cell(table.rows[0].cells[ci], h, '16a085')
adv_data = [
    ['Advantage (เสริมแกร่ง)', 'ทอย Attack Roll (2d6) 2 ครั้ง เอาค่าที่ มากกว่า'],
    ['Disadvantage (มึนงง)', 'ทอย Attack Roll (2d6) 2 ครั้ง เอาค่าที่ น้อยกว่า'],
]
for ri, row_data in enumerate(adv_data):
    fill = 'd5f5e3'
    for ci, text in enumerate(row_data):
        cell = table.rows[ri+1].cells[ci]
        set_cell_shading(cell, fill)
        set_cell_text(cell, text, size=Pt(9), font_name='Arial Unicode MS', color='333333')

add_normal(doc, 'Advantage ได้จาก: สกิล Empower ของ Priest / ไอเทม Elixir of Power')
add_normal(doc, 'Disadvantage ได้จาก: สกิล Daze ของ Priest / สกิลพิเศษบางมอนฯ (เช่น Slime)')

doc.add_paragraph(' ')

# 7.3 กฎการโจมตีของมอนสเตอร์
add_heading2(doc, '7.3 กฎการโจมตีของมอนสเตอร์')
add_normal(doc, 'มอนสเตอร์ต้องทอย Attack Roll (2d6) เหมือนผู้เล่น ระบบ: ทอยได้ 2 = Critical Failure (พลาด 0 damage) | ทอยได้ < AC = Glancing Blow (damage ÷ 2 ปัดลง ขั้นต่ำ 1) | ทอยได้ >= AC = Normal Hit (damage เต็ม) | ทอยได้ 12 = Critical Hit (damage × 2)', bold=True)
add_normal(doc, 'ตัวอย่าง: พี่มอนฯ ทอย 2d6 ได้ 9 เป้าหมายคือ Fighter (AC 7) -> 9 >= 7 ตีโดน! -> ทอย damage ตามใบอธิบาย | ถ้าทอยได้ 5 -> 5 < 7 = Glancing Blow damage ÷ 2')

doc.add_paragraph(' ')

# 7.4 Non-Attack Skills
add_heading2(doc, '7.4 สกิลที่ไม่ใช่การโจมตี (Non-Attack Skills)')
add_normal(doc, 'สกิล Heal, Amplify Damage, Shield Wall, Armor Break, Daze, Empower ไม่ต้องทอย Attack Roll ใช้ได้เลยในเทิร์นของตัวเอง มีผลทันที')

doc.add_paragraph(' ')

# 7.5 การกระทำในแต่ละเทิร์น
add_heading2(doc, '7.5 การกระทำในแต่ละเทิร์น')
add_normal(doc, 'ใช้สกิล 1 อย่าง + ใช้ไอเทม 1 ชิ้น (จะใช้ก่อนหรือหลังสกิลก็ได้) ไอเทมเป็นของส่วนกลางทั้งทีม')

doc.add_paragraph(' ')

# 7.6 กรณีมีคนตาย
add_heading2(doc, '7.6 กรณีมีคนตายระหว่างไฟท์ / TPK')
add_normal(doc, 'มีคนตาย แต่มอนฯ ตายก่อน: ได้รางวัลเหมือนเดิม คนที่ตายเหลือเลือด 1 HP เพิ่มไม่ได้จนกว่าจะฮีล')
# UPDATED TPK text
add_normal(doc, 'TPK (ตายหมด): เสีย LOC ทั้งหมด เสียไอเทมทั้งหมด สกิลถาวร (สกิล class + Scroll ที่ปลดล็อคแล้ว) + Passive ยังอยู่ สกิลใช้ครั้งเดียว (เช่น Stun Strike ที่ยังไม่ได้ใช้) จะหายไป Full Heal ที่ร้านค้าหลัง TPK = ฟรี (0 LOC)')

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# ==================== 8. ประเภท Damage & ระบบจุดอ่อน ====================
add_heading1(doc, '8. ประเภท Damage & ระบบจุดอ่อน')

# 8.1 Damage ทั้ง 6 ประเภท
add_heading2(doc, '8.1 Damage ทั้ง 6 ประเภท')

doc.add_paragraph(' ')

# Damage type table (Table 7)
table = doc.add_table(rows=7, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(['ประเภท Damage', 'คำอธิบาย']):
    make_header_cell(table.rows[0].cells[ci], h, '8e44ad')
dmg_data = [
    ['Physical \U0001f5e1\ufe0f', 'ดาเมจจากดาบ, คทา, อาวุธทั่วไป มอนฯ หลายตัว Resist ได้'],
    ['Fire \U0001f525', 'ดาเมจธาตุไฟ ได้ผลดีกับ Undead และมอนฯ ธีมน้ำแข็ง'],
    ['Frost \u2744\ufe0f', 'ดาเมจธาตุน้ำแข็ง ได้ผลดีกับมอนฯ ธีมไฟ'],
    ['Lightning \u26a1', 'ดาเมจธาตุสายฟ้า ได้ผลดีกับมอนฯ สวมเกราะโลหะ'],
    ['Necrotic \U0001f480', 'ดาเมจพลังมืด Undead มักจะ Immune'],
    ['Holy \U0001f31f', 'ดาเมจพลังศักดิ์สิทธิ์ ได้ผลดีมากกับ Undead และปีศาจ'],
]
for ri, row_data in enumerate(dmg_data):
    for ci, text in enumerate(row_data):
        make_data_cell(table.rows[ri+1].cells[ci], text)

# 8.2 Vulnerability System
add_heading2(doc, '8.2 ระบบจุดอ่อน (Vulnerability System)')

doc.add_paragraph(' ')

# Vuln table (Table 8)
table = doc.add_table(rows=4, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(['สถานะ', 'ผลต่อ Damage', 'คำอธิบาย']):
    make_header_cell(table.rows[0].cells[ci], h, 'c0392b')
vuln_data = [
    ['Vulnerable (แพ้ทาง)', 'Damage x2', 'ดาเมจประเภทนี้ทำ damage เป็น 2 เท่า'],
    ['Resistant (ต้านทาง)', 'Damage /2', 'ดาเมจประเภทนี้ลดลงครึ่งหนึ่ง (ปัดลง)'],
    ['Immune (ชนะทาง)', 'Damage = 0', 'ดาเมจประเภทนี้ไม่มีผลเลย'],
]
for ri, row_data in enumerate(vuln_data):
    fill = 'fadbd8'
    for ci, text in enumerate(row_data):
        cell = table.rows[ri+1].cells[ci]
        set_cell_shading(cell, fill)
        set_cell_text(cell, text, size=Pt(9), font_name='Arial Unicode MS', color='333333')

add_normal(doc, 'สำคัญ: น้อง ๆ จะไม่รู้จุดอ่อนของมอนฯ ต้องสังเกตจากลักษณะภายนอกเอง เช่น ปีศาจเกราะเหล็ก น่าจะแพ้ Holy/Lightning แต่ต้านทาน Physical', italic=True)

doc.add_paragraph(' ')

# ==================== 9. แผนที่ดันเจี้ยน ====================
add_heading1(doc, '9. แผนที่ดันเจี้ยน (Dungeon Map)')
add_normal(doc, 'ดันเจี้ยนกระจายอยู่ตามจุดต่าง ๆ ของมหาวิทยาลัย รุ่นพี่นั่งประจำแต่ละจุด ทีมต้องเดินไปหา')

doc.add_paragraph(' ')
doc.add_paragraph(' ')

# Map table (Table 9)
table = doc.add_table(rows=8, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(['สถานที่', 'จำนวน', 'รายละเอียด']):
    make_header_cell(table.rows[0].cells[ci], h, 'd35400')
map_data = [
    ['มอนสเตอร์ Rank C', '3', 'มอนฯ อ่อนสุด เหมาะฟาร์มเงินและฝึกมือ | Cooldown: 2 นาที'],
    ['มอนสเตอร์ Rank B', '2', 'ความยากปานกลาง ดรอปของดีกว่า | Cooldown: 3 นาที'],
    ['มอนสเตอร์ Rank A', '1', 'มอนฯ แข็งแกร่ง ดรอปมีค่ามาก | Cooldown: 5 นาที'],
    ['มอนสเตอร์ Rank S (Duo)', '1 จุด (2 ตัว)', 'มอนฯ คู่หู ต้องรับมือพร้อมกัน | Cooldown: 10 นาที'],
    ['Boss (Solo)', '1', 'บอสเดี่ยว ไม่มีวันตาย ตีได้หลายรอบ บันทึก Total Damage'],
    ['ร้านค้า', '1', 'ร้านค้ารวม ไอเทมทุกชนิด + Full Heal'],
    ['-', '-', '(ร้านค้ารวมเป็นร้านเดียว)'],
]
for ri, row_data in enumerate(map_data):
    for ci, text in enumerate(row_data):
        make_data_cell(table.rows[ri+1].cells[ci], text)

add_normal(doc, 'กฎแผนที่: ห้ามตีมอนฯ ตัวเดิมติดกัน ต้องไปตีตัวอื่นก่อน 1 ตัว แล้วรอ Spawn Cooldown ถึงจะกลับมาตีได้', bold=True)
add_normal(doc, '[TBD] แนบแผนที่มหาวิทยาลัยพร้อมจุดมอนฯ/ร้านค้า/บอส ก่อนวันงาน', italic=True)

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# ==================== 10. มอนสเตอร์ทั้งหมด ====================
add_heading1(doc, '10. มอนสเตอร์ทั้งหมด')

# 10.1 Rank C
add_heading2(doc, '10.1 Rank C - มอนฯ ระดับต้น (3 ตัว)')
add_normal(doc, 'Rank C ไม่มี Vulnerable/Resistant/Immune เพื่อให้ง่ายสำหรับน้อง ๆ ที่เพิ่งเริ่มเล่น', italic=True)

doc.add_paragraph(' ')

# C1: Goblin Scout - UPDATED: HP 18 (was 14), target last_attacker
add_heading3(doc, 'C1: Goblin Scout (ก็อบลินลาดตระเวน)')

doc.add_paragraph(' ')

add_monster_table(doc, 'C1', '27ae60', [
    ['สเตตัส', 'Goblin Scout (ก็อบลินลาดตระเวน) (C)'],
    ['ชื่อ', 'Goblin Scout (ก็อบลินลาดตระเวน)'],
    ['Rank', 'C'],
    ['HP', '18'],
    ['AC', '5'],
    ['Spawn Cooldown', '2 นาที'],
    ['รูปแบบโจมตี', 'โจมตีตัวที่ตีเราล่าสุด: 1d6+1d4 Physical Damage \U0001f5e1\ufe0f'],
    ['ประเภท Damage', 'Physical \U0001f5e1\ufe0f'],
    ['Vulnerable', '-'],
    ['Resistant', '-'],
    ['Immune', '-'],
    ['ของดรอป', '15 LOC + Potion of Healing (50%)'],
    ['กิมมิคพิเศษ', 'เคลื่อนที่เร็ว แต่อ่อนแอ เหมาะฟาร์มเงิน'],
])

doc.add_paragraph(' ')

# C2: Skeleton Warrior - UPDATED: HP 20 (was 16), target init_cycle
add_heading3(doc, 'C2: Skeleton Warrior (โครงกระดูกนักรบ)')

doc.add_paragraph(' ')

add_monster_table(doc, 'C2', '27ae60', [
    ['สเตตัส', 'Skeleton Warrior (โครงกระดูกนักรบ) (C)'],
    ['ชื่อ', 'Skeleton Warrior (โครงกระดูกนักรบ)'],
    ['Rank', 'C'],
    ['HP', '20'],
    ['AC', '6'],
    ['Spawn Cooldown', '2 นาที'],
    ['รูปแบบโจมตี', 'โจมตีตามลำดับ Initiative (วนรอบ): 2d6 Physical Damage \U0001f5e1\ufe0f\nทุก 2 เทิร์น: AoE ตีทั้งทีม 1d6+1d4 Physical Damage \U0001f5e1\ufe0f'],
    ['ประเภท Damage', 'Physical \U0001f5e1\ufe0f'],
    ['Vulnerable', '-'],
    ['Resistant', '-'],
    ['Immune', 'Necrotic \U0001f480'],
    ['ของดรอป', '20 LOC + Scroll: Frost Bolt (40%) + Scroll: Daze (30%) + Scroll: Armor Break (10%)'],
    ['กิมมิคพิเศษ', 'Undead - สกิล Heal ของ Priest ถ้าใช้กับมอนฯ ตัวนี้ จะทำ damage x2 แทน'],
])

doc.add_paragraph(' ')

# C3: Slime - UPDATED: HP 16 (was 12), target random
add_heading3(doc, 'C3: Slime (สไลม์เหนียว)')

doc.add_paragraph(' ')

add_monster_table(doc, 'C3', '27ae60', [
    ['สเตตัส', 'Slime (สไลม์เหนียว) (C)'],
    ['ชื่อ', 'Slime (สไลม์เหนียว)'],
    ['Rank', 'C'],
    ['HP', '16'],
    ['AC', '5'],
    ['Spawn Cooldown', '2 นาที'],
    ['รูปแบบโจมตี', 'โจมตีสุ่มเป้าหมาย: 1d6+1d4 Physical Damage \U0001f5e1\ufe0f + ติด Disadvantage 1 เทิร์น'],
    ['ประเภท Damage', 'Physical \U0001f5e1\ufe0f'],
    ['Vulnerable', '-'],
    ['Resistant', '-'],
    ['Immune', '-'],
    ['ของดรอป', '20 LOC + Random Scroll (25%)'],
    ['กิมมิคพิเศษ', 'ตีไม่แรง แต่ทำให้มึนงง (Disadvantage) เมื่อ Slime ถูกฆ่าครบ 3 ครั้ง (รวมจากทุกทีม) จะกลายเป็น Slime King (Rank A)'],
])

doc.add_paragraph(' ')

# C3(A): Slime King - UPDATED: HP 16 (was 48), AC 10 (was 12), target init_cycle
add_heading3(doc, 'C3 (A): Slime King (สไลม์คิง) - Hidden Evolution')

doc.add_paragraph(' ')

add_monster_table(doc, 'C3A', 'e74c3c', [
    ['สเตตัส', 'Slime King (สไลม์คิง) (A)'],
    ['ชื่อ', 'Slime King (สไลม์คิง)'],
    ['Rank', 'A (วิวัฒนาการจาก Slime หลังตายครบ 3 ครั้งรวมทุกทีม)'],
    ['HP', '16'],
    ['AC', '10'],
    ['Spawn Cooldown', '2 นาที (กลับเป็น Slime ธรรมดาเมื่อถูกฆ่า)'],
    ['รูปแบบโจมตี', 'โจมตีตามลำดับ Initiative (วนรอบ): 2d6 Physical Damage \U0001f5e1\ufe0f + Disadvantage 1 เทิร์น'],
    ['ประเภท Damage', 'Physical \U0001f5e1\ufe0f'],
    ['Vulnerable', 'Lightning \u26a1'],
    ['Resistant', 'Physical \U0001f5e1\ufe0f'],
    ['Immune', 'Frost \u2744\ufe0f'],
    ['ของดรอป', '70 LOC + Random Item + Random Scroll (50%)'],
    ['กิมมิคพิเศษ', 'เมื่อถูกฆ่าจะกลับไปเป็น Slime ธรรมดา นับ kill ใหม่'],
])

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# 10.2 Rank B
add_heading2(doc, '10.2 Rank B - มอนฯ ระดับกลาง (2 ตัว)')

doc.add_paragraph(' ')

# B1: Dark Knight - UPDATED: HP 13 (was 20), AC 8 (was 6), attack 2d6 (not 2d6+1d4), NO AOE
add_heading3(doc, 'B1: Dark Knight (อัศวินทมิฬ)')

doc.add_paragraph(' ')

add_monster_table(doc, 'B1', '2980b9', [
    ['สเตตัส', 'Dark Knight (อัศวินทมิฬ) (B)'],
    ['ชื่อ', 'Dark Knight (อัศวินทมิฬ)'],
    ['Rank', 'B'],
    ['HP', '13'],
    ['AC', '8'],
    ['Spawn Cooldown', '3 นาที'],
    ['รูปแบบโจมตี', 'โจมตีตัวที่ตีเราล่าสุด: 2d6 Physical Damage \U0001f5e1\ufe0f'],
    ['ประเภท Damage', 'Physical \U0001f5e1\ufe0f'],
    ['Vulnerable', 'Holy \U0001f31f, Lightning \u26a1'],
    ['Resistant', 'Physical \U0001f5e1\ufe0f'],
    ['Immune', 'Necrotic \U0001f480'],
    ['ของดรอป', '40 LOC + Iron Ring (100%) + Scroll: Lightning (40%) + Scroll: Armor Break (20%)'],
    ['กิมมิคพิเศษ', 'เกราะหนัก AC สูง ต้องใช้ Armor Break หรือ Elemental Damage ถึงจะมีประสิทธิภาพ'],
])

doc.add_paragraph(' ')

# B2: Flame Serpent - UPDATED: HP 40 (was 28), AC 6 (was 8), attack 2d6+1d4 Fire (target last_attacker), heal 1d6 every 3 turns
add_heading3(doc, 'B2: Flame Serpent (งูเพลิง)')

doc.add_paragraph(' ')

add_monster_table(doc, 'B2', '2980b9', [
    ['สเตตัส', 'Flame Serpent (งูเพลิง) (B)'],
    ['ชื่อ', 'Flame Serpent (งูเพลิง)'],
    ['Rank', 'B'],
    ['HP', '40'],
    ['AC', '6'],
    ['Spawn Cooldown', '3 นาที'],
    ['รูปแบบโจมตี', 'ปกติ: โจมตีตัวที่ตีเราล่าสุด 2d6+1d4 Fire Damage \U0001f525\nทุก 3 เทิร์น: ฮีลตัวเอง 1d6'],
    ['ประเภท Damage', 'Fire \U0001f525'],
    ['Vulnerable', 'Frost \u2744\ufe0f'],
    ['Resistant', '-'],
    ['Immune', 'Fire \U0001f525 (จะฮีลถ้าโดนตีด้วยไฟ)'],
    ['ของดรอป', '45 LOC + Scroll: Necrotic Blast (30%) + Potion of Healing (100%) + Scroll: Empower (30%)'],
    ['กิมมิคพิเศษ', 'ตี Fire Damage เป็นหลัก ฮีลตัวเองทุก 3 เทิร์น ใช้ Frost Bolt ได้ผลดีมาก'],
])

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# 10.3 Rank A
add_heading2(doc, '10.3 Rank A - มอนฯ ระดับสูง (1 ตัว)')

doc.add_paragraph(' ')

# A: Lich King - UPDATED: HP 44 (was 24), attack 2d6+1d4 (target highest_hp), drain every 3 turns 2d6, Phase 2 12 HP + 3d6, Vuln Holy ONLY (not Fire+Holy)
add_heading3(doc, 'A: Lich King (พ่อมดมรณะ)')

doc.add_paragraph(' ')

add_monster_table(doc, 'A', 'e74c3c', [
    ['สเตตัส', 'Lich King (พ่อมดมรณะ) (A)'],
    ['ชื่อ', 'Lich King (พ่อมดมรณะ)'],
    ['Rank', 'A'],
    ['HP', '44'],
    ['AC', '7'],
    ['Spawn Cooldown', '5 นาที'],
    ['รูปแบบโจมตี (Phase 1)', 'ปกติ: โจมตีตัวที่ HP มากสุด 2d6+1d4 Necrotic Damage \U0001f480\nทุก 3 เทิร์น: ดูด HP 2d6 Necrotic (ฮีลตัวเองเท่ากับ damage ที่ทำได้)'],
    ['Phase 2 (หลังตาย)', 'ทำหน้าเหมือนตาย แล้วฟื้นด้วย 12 HP\nปกติ: โจมตีตัวที่ HP มากสุด 3d6 Necrotic Damage \U0001f480\nทุก 3 เทิร์น: ดูด HP 2d6 Necrotic (ฮีลตัวเอง)'],
    ['ประเภท Damage', 'Necrotic \U0001f480'],
    ['Vulnerable', 'Holy \U0001f31f'],
    ['Resistant', 'Frost \u2744\ufe0f'],
    ['Immune', 'Necrotic \U0001f480'],
    ['ของดรอป', '65 LOC + Lich Crown (100%) + Scroll: Taunt (50%)'],
    ['กิมมิคพิเศษ', 'ดูด HP ฮีลตัวเอง ต้องรีบจัดการ ใช้ Holy ได้ผลดี Phase 2 แรงขึ้นมาก ต้องจัดการให้เร็ว'],
])

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# 10.4 Rank S - Duo
add_heading2(doc, '10.4 Rank S - มอนฯ ระดับสูงมาก (Duo)')

doc.add_paragraph(' ')

add_heading3(doc, 'S: Queen Divine & King Conquer (ราชินีแดนสรวง & ราชาผู้พิชิต)')

doc.add_paragraph(' ')

# Queen/King Duo table - UPDATED: Queen HP 13 (was 16), Queen attack 1d6 Holy (was 2d6), King Resist Physical (was Frost/Necrotic/Fire), King Immune -
duo_data = [
    ['สเตตัส', 'Queen Divine & King Conquer (Rank S - Duo)'],
    ['Rank', 'S (Duo) - 2 ตัวสู้พร้อมกัน'],
    ['Spawn Cooldown', '10 นาที'],
    ['--- ตัวที่ 1 ---', '--- Queen Divine (ราชินีแดนสรวง) ---'],
    ['HP ตัวที่ 1', '13'],
    ['AC ตัวที่ 1', '4'],
    ['โจมตีตัวที่ 1', 'ปกติ: โจมตีตัวที่ตีเราล่าสุด 1d6 Holy Damage \U0001f31f\nทุก 2 เทิร์น: Disadvantage ทั้งทีม 1 เทิร์น\nทุก 3 เทิร์น: สร้างเกราะให้ King Conquer ถาวร +1 AC'],
    ['Vulnerable ตัวที่ 1', 'Necrotic \U0001f480'],
    ['Resistant ตัวที่ 1', 'Frost \u2744\ufe0f, Physical \U0001f5e1\ufe0f'],
    ['Immune ตัวที่ 1', 'Holy \U0001f31f'],
    ['--- ตัวที่ 2 ---', '--- King Conquer (ราชาผู้พิชิต) ---'],
    ['HP ตัวที่ 2', '24'],
    ['AC ตัวที่ 2', '7'],
    ['โจมตีตัวที่ 2', 'ปกติ: โจมตีตัวที่ HP น้อยสุด 2d6 Physical Damage \U0001f5e1\ufe0f\nทุก 3 เทิร์น: ตีทั้งทีม 2d6 Physical Damage \U0001f5e1\ufe0f (ทอยแยกรายคน)'],
    ['Vulnerable ตัวที่ 2', 'Holy \U0001f31f, Lightning \u26a1'],
    ['Resistant ตัวที่ 2', 'Physical \U0001f5e1\ufe0f'],
    ['Immune ตัวที่ 2', '-'],
    ['ของดรอป', '100 LOC + Scroll: Holy Light + Phoenix Feather'],
    ['กิมมิคพิเศษ Duo', 'ทั้ง 2 ตัวทอย Initiative แยกกัน (มีเทิร์นคนละเวลา) Queen buff King ด้วย +AC ทำให้ยิ่งสู้นาน King ยิ่งแข็ง Queen Immune ต่อ Holy แต่แพ้ Necrotic King แพ้ Holy/Lightning -> ใช้คนละธาตุจัดการ'],
]

table = doc.add_table(rows=len(duo_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_cell(table.rows[0].cells[0], duo_data[0][0], '8e44ad')
make_header_cell(table.rows[0].cells[1], duo_data[0][1], '8e44ad')
for i in range(1, len(duo_data)):
    left_text, right_text = duo_data[i]
    fill = 'f8f9fa'
    # Separator rows
    if '---' in left_text:
        fill = 'e8daef'
    set_cell_shading(table.rows[i].cells[0], fill)
    set_cell_text(table.rows[i].cells[0], left_text, bold=True, size=Pt(9), font_name='Arial Unicode MS', color='333333')
    set_cell_shading(table.rows[i].cells[1], fill)
    if '\n' in right_text:
        lines = right_text.split('\n')
        set_cell_multiline(table.rows[i].cells[1], lines, bold_first=False)
    else:
        set_cell_text(table.rows[i].cells[1], right_text, size=Pt(9), font_name='Arial Unicode MS', color='333333')

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# ==================== 11. บอส (Boss) ====================
add_heading1(doc, '11. บอส (Boss)')
add_normal(doc, 'บอสไม่มีวันตาย พี่ที่ประจำจุดจดบันทึก Total Damage ที่แต่ละทีมทำได้ ทีมสามารถ Raid บอสได้หลายรอบไม่จำกัด Damage สะสมตลอด! Damage จาก Rank S (Duo) ก็นับรวมด้วย!', bold=True)

doc.add_paragraph(' ')

# 11.1 Boss - UPDATED: AC 7 (was 8), 5 turns per attempt
add_heading2(doc, '11.1 Boss - Infernal Demon Lord (ราชาปีศาจนรก) [Solo]')

doc.add_paragraph(' ')

boss_data = [
    ['สเตตัส', 'Infernal Demon Lord - ราชาปีศาจนรก (Boss Solo)'],
    ['Rank', 'Boss (Solo)'],
    ['HP', 'ไม่มีวันตาย (จด Total Damage ที่ทีมทำได้)'],
    ['AC', '7'],
    ['รูปแบบโจมตี', 'ปกติ: โจมตีตัวที่ทำ damage สะสมสูงสุด: 3d6 Fire Damage \U0001f525\nทุก 2 เทิร์น: Hellfire Sweep - ตีทั้งทีม 2d6 Fire Damage \U0001f525\nทุก 4 เทิร์น: Infernal Rage - โจมตีตัว AC น้อยสุด 4d6 Fire Damage \U0001f525'],
    ['ประเภท Damage', 'Fire \U0001f525'],
    ['Vulnerable', 'Holy \U0001f31f'],
    ['Resistant', 'Physical \U0001f5e1\ufe0f, Lightning \u26a1'],
    ['Immune', 'Necrotic \U0001f480, Fire \U0001f525'],
    ['กิมมิคพิเศษ', 'บอสแพ้ทาง Holy แน่นอน ทีมที่มี Holy Light จะได้เปรียบมาก\nInfernal Rage ทุก 4 เทิร์น เล็ง AC ต่ำสุด (Mage/Priest) ทำ damage มหาศาล ต้องฮีลให้ทัน\nResistant Cloak จะช่วยลด damage ได้\nBoss Raid: 5 เทิร์นต่อรอบ ตีได้ไม่จำกัดรอบ Damage สะสมตลอด!\nDamage จาก Duo (Rank S) นับรวมเป็น Total Damage ด้วย!'],
]

table = doc.add_table(rows=len(boss_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
make_header_cell(table.rows[0].cells[0], boss_data[0][0], 'c0392b')
make_header_cell(table.rows[0].cells[1], boss_data[0][1], 'c0392b')
for i in range(1, len(boss_data)):
    left_text, right_text = boss_data[i]
    fill = 'f8f9fa'
    set_cell_shading(table.rows[i].cells[0], fill)
    set_cell_text(table.rows[i].cells[0], left_text, bold=True, size=Pt(9), font_name='Arial Unicode MS', color='333333')
    set_cell_shading(table.rows[i].cells[1], fill)
    if '\n' in right_text:
        lines = right_text.split('\n')
        set_cell_multiline(table.rows[i].cells[1], lines, bold_first=False)
    else:
        set_cell_text(table.rows[i].cells[1], right_text, size=Pt(9), font_name='Arial Unicode MS', color='333333')

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# ==================== 12. ไอเทม & ร้านค้า ====================
add_heading1(doc, '12. ไอเทม & ร้านค้า')

# 12.1 สกุลเงิน
add_heading2(doc, '12.1 สกุลเงิน')
add_normal(doc, 'เงินในเกมคือ LOC Coins ได้จากการตีมอนสเตอร์')

doc.add_paragraph(' ')

# 12.2 รายการไอเทมทั้งหมด - UPDATED prices and effects
add_heading2(doc, '12.2 รายการไอเทมทั้งหมด')

doc.add_paragraph(' ')

# Shop table (Table 19) - UPDATED
shop_items = [
    ['ชื่อไอเทม', 'ประเภท', 'ราคา', 'ผลลัพธ์', 'ขายที่'],
    ['Potion of Healing', 'Consumable', '15 LOC', 'ฮีล HP ตัวละคร 1 ตัว = 1d6+1d4 HP', 'ร้านค้า'],
    ['Full Heal', 'บริการ', '30 LOC (0 หาก TPK)', 'ฮีล HP ทั้งทีมกลับเต็ม (ทุกตัว = Max HP)', 'ร้านค้า'],
    ['Iron Ring', 'Passive', '35 LOC', 'เพิ่ม AC ทั้งทีมขึ้น 1 ตลอดเกม', 'ร้านค้า'],
    ['Resistant Cloak', 'Passive', '35 LOC', 'ลด Damage ที่ได้รับทุกประเภทลง 1 ตลอดเกม (ทั้งทีม)', 'ร้านค้า'],
    ['Scroll: Taunt', 'Skill Unlock', '40 LOC', 'ปลดล็อค Taunt ให้ Fighter (ถาวร)', 'ร้านค้า'],
    ['Scroll: Second Wind', 'Skill Unlock', '40 LOC', 'ปลดล็อค Second Wind ให้ Fighter (ถาวร)', 'ร้านค้า'],
    ['Scroll: Frost Bolt', 'Skill Unlock', '30 LOC', 'ปลดล็อค Frost Bolt ให้ Mage (ถาวร)', 'ร้านค้า'],
    ['Scroll: Lightning', 'Skill Unlock', '40 LOC', 'ปลดล็อค Lightning ให้ Mage (ถาวร)', 'ร้านค้า'],
    ['Scroll: Necrotic Blast', 'Skill Unlock', '40 LOC', 'ปลดล็อค Necrotic Blast ให้ Mage (ถาวร)', 'ร้านค้า'],
    ['Scroll: Armor Break', 'Skill Unlock', '40 LOC', 'ปลดล็อค Armor Break ให้ Priest (ถาวร)', 'ร้านค้า'],
    ['Scroll: Daze', 'Skill Unlock', '40 LOC', 'ปลดล็อค Daze ให้ Priest (ถาวร)', 'ร้านค้า'],
    ['Scroll: Empower', 'Skill Unlock', '40 LOC', 'ปลดล็อค Empower ให้ Priest (ถาวร)', 'ร้านค้า'],
    ['Scroll: Stun Strike', 'Temporary', '20 LOC', 'ใช้ Stun Strike 1 ครั้ง (หายหลังใช้)', 'ร้านค้า'],
    ['Elixir of Power', 'Consumable', '20 LOC', 'ตัวละคร 1 ตัว ได้ Advantage 2 เทิร์น', 'ร้านค้า'],
    ['Phoenix Feather', 'Consumable', '40 LOC', 'ชุบชีวิตตัวละครที่ตาย กลับมา HP เต็ม (ใช้ได้ทันที)', 'ร้านค้า'],
    ['Holy Water', 'Consumable', '15 LOC', 'ทำ Holy Damage 2d6 ให้มอนฯ 1 ตัว (ใครก็ใช้ได้)', 'ร้านค้า'],
    ['Whetstone', 'Consumable', '15 LOC', 'เพิ่ม damage ครั้งถัดไป +3 (Fighter เท่านั้น)', 'ร้านค้า'],
    ['Stun Bomb', 'Consumable', '30 LOC', 'Stun มอนฯ 1 เทิร์น (ใครก็ใช้ได้)', 'ร้านค้า'],
]

table = doc.add_table(rows=len(shop_items), cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(shop_items[0]):
    make_header_cell(table.rows[0].cells[ci], h, '2980b9')
for ri in range(1, len(shop_items)):
    for ci, text in enumerate(shop_items[ri]):
        make_data_cell(table.rows[ri].cells[ci], text)

# 12.3 ของดรอปจากมอนสเตอร์
add_heading2(doc, '12.3 ของดรอปจากมอนสเตอร์')

doc.add_paragraph(' ')

# Drop table (Table 20) - UPDATED
drop_data = [
    ['แหล่ง', 'เงิน LOC', 'ไอเทม/สกิล'],
    ['มอนฯ Rank C', '15-20 LOC', 'Potion of Healing (50%) / Scroll สกิล (Frost Bolt 40%, Daze 30%, Armor Break 10%) / Random Scroll (25%)'],
    ['มอนฯ Rank B', '40-45 LOC', 'Iron Ring (100%) / Scroll: Lightning (40%) / Scroll: Armor Break (20%) / Necrotic Blast (30%) / Potion of Healing (100%) / Scroll: Empower (30%)'],
    ['มอนฯ Rank A', '65 LOC', 'Lich Crown (100%) + Scroll: Taunt (50%)'],
    ['มอนฯ Rank S', '100 LOC', 'Scroll: Holy Light + Phoenix Feather'],
    ['Boss', 'ไม่ดรอป', 'บอสไม่มีวันตาย เป้าหมายคือทำ damage ให้มากที่สุด'],
]

table = doc.add_table(rows=len(drop_data), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(drop_data[0]):
    make_header_cell(table.rows[0].cells[ci], h, 'f39c12')
for ri in range(1, len(drop_data)):
    for ci, text in enumerate(drop_data[ri]):
        make_data_cell(table.rows[ri].cells[ci], text)

add_normal(doc, 'หมายเหตุ: เปอร์เซ็นต์ดรอปอาจปรับตามดุลยพินิจของพี่ประจำจุด หรือ DM สามารถทอยลูกเต๋าตัดสิน', italic=True)

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# ==================== 13. กฎการตาย & การแพ้ ====================
add_heading1(doc, '13. กฎการตาย & การแพ้')

doc.add_paragraph(' ')

# TPK table (Table 21) - UPDATED
tpk_data = [
    ['สถานการณ์', 'ผลลัพธ์'],
    ['ตัวละครตาย 1 ตัว (ระหว่างไฟท์)', 'ไฟท์ต่อได้ ตัวที่ตาย skip เทิร์น เมื่อจบไฟท์ ฟื้นกลับมาเป็น 1 HP ฮีลระหว่างทางไม่ได้ ยกเว้นใช้ Potion / Phoenix Feather / Full Heal ที่ร้าน'],
    ['ทั้งทีม TPK (ตายหมด)', 'เสีย LOC ทั้งหมด + เสียไอเทมทั้งหมด สกิลถาวร (สกิล class + Scroll ที่ปลดล็อคแล้ว) + Passive ยังอยู่ สกิลใช้ครั้งเดียว (เช่น Stun Strike ที่ยังไม่ได้ใช้) จะหายไป Full Heal ที่ร้านค้าหลัง TPK = ฟรี (0 LOC)'],
    ['หนีจากการต่อสู้', 'ทุกตัวละครถูกมอนฯ ตีคนละ 1 ครั้ง ถ้ารอดชีวิต = หนีสำเร็จ ถ้าตาย = ถือว่า TPK ตามกฎข้างบน'],
]

table = doc.add_table(rows=len(tpk_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(tpk_data[0]):
    make_header_cell(table.rows[0].cells[ci], h, 'c0392b')
for ri in range(1, len(tpk_data)):
    for ci, text in enumerate(tpk_data[ri]):
        make_data_cell(table.rows[ri].cells[ci], text)

# ==================== 14. สรุปกฎสำคัญ (Quick Reference) ====================
add_heading1(doc, '14. สรุปกฎสำคัญ (Quick Reference)')

doc.add_paragraph(' ')

# Quick ref table (Table 22) - UPDATED
qref_data = [
    ['กฎ', 'คำอธิบาย'],
    ['ห้ามตีมอนฯ ซ้ำติดกัน', 'ต้องไปตีตัวอื่นก่อน 1 ตัว ถึงจะกลับมาตีซ้ำได้ (ต้องรอ Spawn Cooldown ด้วย)'],
    ['Spawn Cooldown', 'Rank C: 2 นาที / Rank B: 3 นาที / Rank A: 5 นาที / Rank S: 10 นาที'],
    ['ไอเทมแชร์กันทั้งทีม', 'ไอเทมทุกชิ้นเป็นของทั้งทีม ใครจะใช้ก็ได้'],
    ['สกิลถาวร vs ใช้ครั้งเดียว', 'สกิลพื้นฐาน + Scroll ที่ปลดล็อค = ถาวร / Stun Strike ฯลฯ = ใช้ครั้งเดียว ต้องหาใหม่'],
    ['มอนฯ ต้อง Roll Attack', 'มอนฯ ทอย 2d6 เหมือนผู้เล่น: 2=miss, <AC=dmg÷2, >=AC=full hit, 12=crit×2'],
    ['ฮีลระหว่างทาง', 'ไม่ได้! ฮีลได้เฉพาะ: สกิล Priest ในไฟท์ / Potion ก่อนสู้ได้ / Full Heal ที่ร้าน'],
    ['ลำดับ Turn Action', 'ใช้สกิล 1 + ไอเทม 1 ก่อนหรือหลังก็ได้'],
    ['จุดอ่อนมอนฯ (ลับ!)', 'น้อง ๆ ไม่รู้จุดอ่อนมอนฯ ต้องสังเกตจากลักษณะ/ธีมเอง'],
    ['บอส Raid', 'ตีบอสได้หลายรอบไม่จำกัด 5 เทิร์นต่อรอบ Damage สะสมตลอด'],
    ['Slime -> Slime King', 'เมื่อ Slime ถูกฆ่าครบ 3 ครั้ง (รวมจากทุกทีม) จะกลายเป็น Slime King (Rank A)'],
    ['เงื่อนไขชนะ', 'ทีมที่ทำ Total Damage สูงสุดจาก Duo (Rank S) + บอส ชนะกิจกรรม'],
]

table = doc.add_table(rows=len(qref_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for ci, h in enumerate(qref_data[0]):
    make_header_cell(table.rows[0].cells[ci], h, '2c3e50')
for ri in range(1, len(qref_data)):
    for ci, text in enumerate(qref_data[ri]):
        make_data_cell(table.rows[ri].cells[ci], text)

# ==================== 15. เงื่อนไขชนะ ====================
add_heading1(doc, '15. เงื่อนไขชนะ')
add_normal(doc, 'ทีมที่ทำ Total Damage รวมจาก Rank S (Duo) + บอส สูงที่สุด เป็นผู้ชนะกิจกรรม!', bold=True)

doc.add_paragraph(' ')

add_normal(doc, '\u2022   \tRaid ได้ไม่จำกัดรอบ: ทีมสามารถกลับมาตีบอสกี่ครั้งก็ได้ Damage สะสมตลอด (รวม Duo + Boss)')
p = doc.add_paragraph()
run = p.add_run('\u2022   \t')
run.font.size = Pt(10)
run.font.name = 'Arial Unicode MS'
run.font.color.rgb = RGBColor.from_string('333333')
run = p.add_run('วางแผนดี = ชนะ: ')
run.bold = True
run.font.size = Pt(10)
run.font.name = 'Arial Unicode MS'
run.font.color.rgb = RGBColor.from_string('333333')
run = p.add_run('ทีมที่ฟาร์มของเก่ง มีไอเทมดี รู้จุดอ่อนบอส จะทำ damage ได้มากกว่า')
run.font.size = Pt(10)
run.font.name = 'Arial Unicode MS'
run.font.color.rgb = RGBColor.from_string('333333')

doc.add_paragraph(' ')
doc.add_paragraph()
doc.add_paragraph(' ')

# ==================== 16. บันทึก Balance สำหรับทีมจัดงาน ====================
add_heading1(doc, '16. บันทึก Balance สำหรับทีมจัดงาน')
add_normal(doc, 'ส่วนนี้สำหรับทีมจัดงานใช้อ้างอิงในการปรับ balance', italic=True)

doc.add_paragraph(' ')

# 16.1
add_heading2(doc, '16.1 หลักการออกแบบ')
add_normal(doc, '\u2022   \tRank C ควรจบได้ใน 3-5 เทิร์น (ฝึกมือน้อง ๆ)')
add_normal(doc, '\u2022   \tRank B ควรจบได้ใน 5-8 เทิร์น (ต้องใช้ strategy บ้าง)')
add_normal(doc, '\u2022   \tRank A ควรจบได้ใน 8-12 เทิร์น (ยากจริง ต้องวางแผน)')
add_normal(doc, '\u2022   \tBoss ตีได้ 5 เทิร์นต่อรอบ ตีได้หลายรอบ Damage สะสม')

doc.add_paragraph(' ')

# 16.2
add_heading2(doc, '16.2 Economy Balance')
add_normal(doc, '\u2022   \tทีมควรตี Rank C 2-3 ตัว ถึงจะมีเงินซื้อ Full Heal ได้ 1 ครั้ง (30 LOC)')
add_normal(doc, '\u2022   \tRank S ดรอป Holy Light ซึ่งสำคัญมากสำหรับตีบอส (Infernal Demon Lord แพ้ทาง Holy)')
add_normal(doc, '\u2022   \tร้านค้ามีทุกอย่างรวมกัน: Potion, Full Heal, Scroll ปลดล็อคสกิล, ไอเทมพิเศษ')
add_normal(doc, '\u2022   \tทีมที่ฟาร์มเก่ง + ใช้เงินฉลาด จะได้เปรียบมาก')

doc.add_paragraph(' ')

# 16.3 - same as spec
add_heading2(doc, '16.3 Damage คาดการณ์ต่อเทิร์น (Average)')
add_normal(doc, '\u2022   \tFighter Sword Slash: avg 5.5/เทิร์น (1d6+1d4)')
add_normal(doc, '\u2022   \tMage Fireball: avg 5.5/เทิร์น (1d6+1d4) | Frost Bolt: avg 5.5 (1d6+1d4) | Lightning/Necrotic: avg 7 (2d6)')
add_normal(doc, '\u2022   \tMage Holy Light: avg 10.5/เทิร์น (3d6) หรือ avg 21 ถ้า Vulnerable!')
add_normal(doc, '\u2022   \tPriest Heal: avg 5.5/เทิร์น (1d6+1d4)')
add_normal(doc, '\u2022   \tPriest Amplify: เพิ่ม avg 3.5/เทิร์น (1d6) ให้เพื่อน')

doc.add_paragraph(' ')

# ==================== 17. สิ่งที่ต้องทำเพิ่มเติม ====================
add_heading1(doc, '17. สิ่งที่ต้องทำเพิ่มเติม')
todo_items = [
    '\u2022   \tแผนที่มหาวิทยาลัยจริงพร้อมจุดสถานีทั้งหมด',
    '\u2022   \tพิมพ์ใบอธิบายมอนสเตอร์สำหรับพี่แต่ละจุด',
    '\u2022   \tพิมพ์ใบติดตามสถานะทีม (Character Sheet) สำหรับ DM',
    '\u2022   \tเตรียมลูกเต๋า d4 และ d6 อย่างน้อย 5 ชุด (ทีมละ 1 ชุด)',
    '\u2022   \tเตรียมเงิน LOC (กระดาษ/เหรียญจำลอง)',
    '\u2022   \tเตรียมการ์ดไอเทม/สกิลสำหรับแจกเมื่อดรอป',
    '\u2022   \tซ้อมพี่ที่เป็นมอนสเตอร์ให้เข้าใจกฎ',
    '\u2022   \tกำหนดตำแหน่งร้านค้าบนแผนที่',
    '\u2022   \tทดลองเล่นจริง 1 รอบเพื่อปรับ balance ก่อนวันงาน',
]
for item in todo_items:
    add_normal(doc, item)

doc.add_paragraph(' ')

# End marker
add_normal(doc, '--- จบเอกสาร ---', italic=True)

doc.add_paragraph()

# ==================== SAVE ====================
output_path = '/Users/vetitk/Coding/LOC/loc-dnd-app/Summer Camp 5 สันเล็ก v4.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print('Done!')
